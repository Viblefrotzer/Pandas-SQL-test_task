# coding=utf-8
# this code is for Python 2.7

import sqlite3
from sqlite3 import Error
import numpy
import pandas
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL

# задаем константой путь к расположению БД
PATH = r"test.db"


# функция коннекта к базе данных
def create_connection(path):
    connection = None
    try:
        connection = sqlite3.connect(path)
        print("Connection to SQLite DB successful")
    except Error:
        print("The error '{}' occurred".format(Error))
    return connection


# создаем курсор и делаем выборку из бд
def execute_read_query(connection, query):
    cursor = connection.cursor()
    result = None
    try:
        cursor.execute(query)
        result = cursor.fetchall()
        return result
    except Error:
        print("The error '{}' occurred".format(Error))


# задаем соединение, используя функцию и путь выше
connect = create_connection(PATH)

# данные выборки из бд для фактора 1
select_data_f1 = """
SELECT * from testidprod
WHERE 
    partner IS NULL 
    AND state IS NULL
    AND bs = 0
    AND factor = 1
"""

# данные выборки из бд для фактора 2
select_data_f2 = """
SELECT * from testidprod
WHERE
    partner IS NULL
    AND state IS NULL
    AND bs = 0
    AND factor = 2
"""

# выполняем два запроса к БД через функцию, получаем данные, закрываем соединение
test_data_f1 = execute_read_query(connection=connect, query=select_data_f1)
test_data_f2 = execute_read_query(connection=connect, query=select_data_f2)
connect.close()

# добавляем названия столбцов для датафрейма
column_names = ['id', 'country', 'factor', 'year', 'res', 'mir', 'raw', 'hash', 'meta', 'partner', 'state', 'bs']
# создаем датафрейм с первым фактором
raw_data_f1 = pandas.DataFrame(data=test_data_f1, columns=column_names)
# исключаем ненужные столбцы для удобства дальнейшей работы
data_f1 = raw_data_f1.drop(labels=["id", "country", "mir", "raw", "hash", "meta", "partner", "state", "bs"],
                           axis=1)
# создаем датафрейм со вторым фактором, выбрасываем ненужные столбцы
raw_data_f2 = pandas.DataFrame(data=test_data_f2, columns=column_names)
data_f2 = raw_data_f2.drop(labels=["id", "country", "mir", "raw", "hash", "meta", "partner", "state", "bs"],
                           axis=1)
# в данных датафреймах отсутствуют данные по 2006 и 2020 годам, т.к. в базе данных нет на них данных
# возможно, есть возможность их добавить автоматом, я её не нашел и добавил нужные столбцы руками.
# Группируя по фактору и году, находим сумму в столбце 'res'
result_f1 = data_f1.groupby(['factor', 'year'])['res'].sum()
result_f2 = data_f2.groupby(['factor', 'year'])['res'].sum()
# переобзываем данные выборки с 'res' на 'world'
result_f1.name = 'world'
result_f2.name = 'world'
# ломаем мультииндекс для транспонирования и дальнейших вычислений
data_f1 = result_f1.reset_index(level=[0, 1])
data_f2 = result_f2.reset_index(level=[0, 1])

# добавляем руками недостающие данные по годам 2006 и 2020 в датафреймы:
nan_data_f1_2006 = [1, 2006, numpy.NAN]
nan_data_f1_2020 = [1, 2020, numpy.NAN]
nan_data_f2_2006 = [2, 2006, numpy.NAN]
nan_data_f2_2020 = [2, 2020, numpy.NAN]
# для датафрейма с фактором 1: data_f1
data_f1.loc[-1] = nan_data_f1_2006  # добавляем ряд
data_f1.index = data_f1.index + 1  # сдвигаем индекс
data_f1.sort_index(inplace=True)  # сортируем, чтобы индекс 0 не был в самом низу
data_f1.loc[14] = nan_data_f1_2020  # добавляем ряд в самый низ
# для датафрейма с фактором 2: data_f2
data_f2.loc[-1] = nan_data_f2_2006  # добавляем ряд
data_f2.index = data_f2.index + 1  # сдвигаем индекс
data_f2.sort_index(inplace=True)  # сортируем, чтобы индекс 0 не был в самом низу
data_f2.loc[14] = nan_data_f2_2020  # добавляем ряд в самый низ
# формируем списки (lists) для датафрейма с фактором 6
# создаем лист с годами, лист с фактором 6
years_list = data_f1['year'].tolist()
# эта конструкция str(int(x)) нужна потому, что у меня не получилось
# при сохранении отчета в word убрать .0 от фактора и года.
# Несмотря на их тип int при сохранении они где-то переводятся во float.
years_list = [str(int(x)) for x in years_list]
new_factor_list = [6 for i in years_list]

# создаем данные столбца world, используя метод zip()
# здесь также можно использовать map()
factor_1_world_data = data_f1['world'].tolist()
factor_2_world_data = data_f2['world'].tolist()
factor_6_world_data = [i / j for i, j in zip(factor_2_world_data, factor_1_world_data)]
# создаем датафрейм c фактором 6
data_f6 = pandas.DataFrame(list(zip(new_factor_list, years_list, factor_6_world_data)),
                           columns=['factor', 'year', 'world'])
# объединяем 3 датафрейма в один
data_all_factors = pandas.concat([data_f1, data_f2, data_f6])
# сбрасываем индексы (появились повторяющиеся после объединения)
data_all_factors = data_all_factors.reset_index()
# удаляем ненужный столбец после предыдущей операции
data_all_factors = data_all_factors.drop(labels=["index"], axis=1)
# переименуем столбцы (в задании в отчете они с большой буквы
data_all_factors = data_all_factors.rename(columns={'factor': 'Factor', 'year': 'Year'})
# сгруппируем по фактору-году, чтобы в экселе получились сгруппированные ячейки на Factor
# транспонируем и сохраняем в эксель
data_all_factors = data_all_factors.set_index(['Factor', 'Year'])
data_all_factors = data_all_factors.transpose()
data_all_factors.to_excel("report.xlsx")

# переименуем столбцы датафрейма с фактором 6
data_f6 = data_f6.rename(columns={'factor': 'Factor', 'year': 'Year', 'world': 'World Value'})
# cоздадим промежуточный лист values для расчетов CAGR. Для это выкинем строки с NaN и сохраним данные
data_f6_for_calc = data_f6.dropna(axis=0)
values = data_f6_for_calc['World Value'].tolist()
years = data_f6_for_calc['Year'].tolist()


# функция для определения CAGR и выдачи строки для репорта
def cagr():
    # заводим вечный цикл для проверки того, что вводимые года есть в нашей таблице
    input_while = True
    while input_while:
        ev = str(input("Enter ending value [year] (f.e. 2019): "))
        bv = str(input("Enter beginning value [year] (f.e. 2007): "))
        if ev < bv:
            print("Invalid data. The ending year value has to be greater then the beginning year value.")
            continue
        if ev not in years or bv not in years:
            print("No data for these years, try again")
            continue
        else:
            # выходим из цикла
            input_while = False
            # определяем значения 'World Value' в выбранные годы
            # по данным строкам мы получаем series с одним значением, поэтому values[0]
            ev_num = data_f6_for_calc['World Value'][data_f6_for_calc['Year'] == ev].values[0]
            bv_num = data_f6_for_calc['World Value'][data_f6_for_calc['Year'] == bv].values[0]
            # формула расчета CAGR: "(EV / BV)^1/n - 1"
            cagr_value = round((((ev_num / bv_num) ** (1 / float(len(values)))) - 1) * 100, 2)
            # проверяем, положительный или отрицательный CAGR получился
            if cagr_value >= 0:
                parameter = 'grew'
            else:
                parameter = 'decreased'
            # и формируем строку отчета
            report_string = "Factor 6 {} by avg {}% every year from {} to {}.".format(parameter, abs(cagr_value), bv,
                                                                                      ev)
            return report_string


# формируем отчет в word
# откроем пустой шаблон doc-файла и будем использовать его для создания отчета
doc = Document()
# создадим таблицу с дополнительным рядом для заголовков
t = doc.add_table(data_f6.shape[0] + 1, data_f6.shape[1])
t.style = 'Table Grid'
# объединим первый столбик таблицы ('Factor')
a = t.cell(1, 0)
b = t.cell(15, 0)
A = a.merge(b)
# расположение значения в середине ячейки
A.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# добавим ряд с заголовками
for j in range(data_f6.shape[-1]):
    t.cell(0, j).text = data_f6.columns[j]
# добавляем остальные данные
for i in range(data_f6.shape[0]):
    for j in range(data_f6.shape[-1]):
        t.cell(i + 1, j).text = str(data_f6.values[i, j])
# добавим строку с расчетом CAGR:
doc.add_paragraph(cagr())
# сохраним документ
doc.save('report.docx')
